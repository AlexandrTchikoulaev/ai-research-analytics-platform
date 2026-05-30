#!/usr/bin/env python3
"""Converte relatorio.txt para relatorio.docx com formatação académica."""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT  = r"relatorio.txt"
OUTPUT = r"relatorio.docx"

RE_H1 = re.compile(r'^(\d+)\.\s{2,}(.+)$')
RE_H2 = re.compile(r'^(\d+\.\d+)\s{2,}(.+)$')
RE_H3 = re.compile(r'^(\d+\.\d+\.\d+)\s{2,}(.+)$')
RE_BULLET = re.compile(r'^- (.+)$')
TABLE_KWORDS = {'Campo', 'Tipo', 'Obrigatório', 'Descrição', 'Etapa', 'Estado', 'Obrigatório'}

DARK_BLUE = RGBColor(0x1F, 0x3A, 0x5F)
MID_GREY  = RGBColor(0x33, 0x33, 0x33)
LITE_GREY = RGBColor(0x55, 0x55, 0x55)


def split_cols(line):
    return [p for p in re.split(r'\s{2,}', line.strip()) if p]


def is_table_header(line):
    parts = split_cols(line)
    return len(parts) >= 2 and any(p.strip() in TABLE_KWORDS for p in parts)


def set_run_font(run, name='Arial', size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold      = bold
    run.italic    = italic
    if color:
        run.font.color.rgb = color


def setup_styles(doc):
    n = doc.styles['Normal']
    n.font.name = 'Arial'
    n.font.size = Pt(11)
    n.paragraph_format.space_after = Pt(6)

    for name, sz, color in [
        ('Heading 1', 16, DARK_BLUE),
        ('Heading 2', 13, DARK_BLUE),
        ('Heading 3', 12, MID_GREY),
        ('Heading 4', 11, LITE_GREY),
    ]:
        s = doc.styles[name]
        s.font.name  = 'Arial'
        s.font.size  = Pt(sz)
        s.font.bold  = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(14)
        s.paragraph_format.space_after  = Pt(4)
        s.paragraph_format.keep_with_next = True

    try:
        lb = doc.styles['List Bullet']
        lb.font.name = 'Arial'
        lb.font.size = Pt(11)
        lb.paragraph_format.space_after = Pt(3)
    except Exception:
        pass


def add_table_block(doc, table_lines):
    rows = [split_cols(l) for l in table_lines if l.strip()]
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < max_cols:
            r.append('')

    tbl = doc.add_table(rows=len(rows), cols=max_cols)
    tbl.style = 'Table Grid'
    for ri, row_data in enumerate(rows):
        cells = tbl.rows[ri].cells
        for ci, text in enumerate(row_data[:max_cols]):
            cell = cells[ci]
            cell.text = ''
            para = cell.paragraphs[0]
            run = para.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.bold = (ri == 0)
    doc.add_paragraph()


def is_subheading(line, prev_blank, next_blank):
    s = line.strip()
    if not s or len(s) > 80:
        return False
    if s.endswith('.') or s.endswith(',') or s.endswith(':'):
        return False
    if RE_H1.match(s) or RE_H2.match(s) or RE_H3.match(s):
        return False
    if RE_BULLET.match(s):
        return False
    # "Etapa N —" or "Cenário N —" are always sub-headings
    if re.match(r'^(Etapa|Cen[aá]rio) \d+', s):
        return True
    # Standalone line (blank both sides) with <= 10 words = sub-heading
    if prev_blank and next_blank:
        if len(s.split()) <= 10:
            return True
    return False


def main():
    with open(INPUT, encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()
    setup_styles(doc)

    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)

    # ── Find boundaries ────────────────────────────────────────────────────
    idx_line   = next((j for j, l in enumerate(lines) if l.strip() == 'Índice'), None)
    # Content starts at the heading immediately before the first long paragraph
    content_start = len(lines)
    if idx_line is not None:
        for j in range(idx_line + 2, len(lines)):
            if len(lines[j].strip()) > 120:
                # Walk back over blank lines to find the heading before this paragraph
                k = j - 1
                while k > idx_line and not lines[k].strip():
                    k -= 1
                if (k > idx_line
                        and lines[k].strip()
                        and (RE_H1.match(lines[k].strip()) or RE_H2.match(lines[k].strip()))):
                    content_start = k
                else:
                    content_start = j
                break

    # ── Cover page ─────────────────────────────────────────────────────────
    i = 0
    cover_end = idx_line if idx_line is not None else 0
    while i < cover_end:
        s = lines[i].strip()
        if s:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(s)
            run.font.name = 'Arial'
            if 'Universidade' in s:
                set_run_font(run, size=15, bold=True)
            elif 'Escola de Ciências' in s:
                set_run_font(run, size=13, bold=True)
            elif 'Projeto em Ciência de Dados' in s:
                set_run_font(run, size=13, bold=True)
            elif re.match(r'^Grupo \d+', s):
                set_run_font(run, size=12, bold=True)
            elif re.match(r'^\w+ \w+ A\d{6}$', s):
                set_run_font(run, size=12)
            elif 'sob a orientação' in s:
                set_run_font(run, size=11, italic=True)
            elif re.match(r'^[A-ZÁÉÍÓÚÂÊÔÃÕ]', s) and len(s) < 50:
                set_run_font(run, size=12)
            else:
                set_run_font(run, size=11)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
        i += 1

    doc.add_page_break()

    # ── Table of Contents ─────────────────────────────────────────────────
    if idx_line is not None:
        doc.add_heading('Índice', level=1)
        i = idx_line + 1
        while i < content_start:
            s = lines[i].strip()
            raw = lines[i]
            if s:
                indent = len(raw) - len(raw.lstrip())
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                if indent >= 4:
                    p.paragraph_format.left_indent = Cm(0.9)
                run = p.add_run(s)
                set_run_font(run, size=10)
            i += 1
        doc.add_page_break()

    # ── Main content ───────────────────────────────────────────────────────
    i = content_start
    n = len(lines)

    while i < n:
        line = lines[i]
        s    = line.strip()

        if not s:
            i += 1
            continue

        prev_blank = (i == 0 or not lines[i-1].strip())
        next_blank = (i+1 >= n or not lines[i+1].strip())

        # Heading 3 first (most specific)
        if RE_H3.match(s):
            doc.add_heading(s, level=3)
            i += 1
            continue

        # Heading 2
        if RE_H2.match(s) and not RE_H3.match(s):
            doc.add_heading(s, level=2)
            i += 1
            continue

        # Heading 1
        if RE_H1.match(s) and not RE_H2.match(s) and not RE_H3.match(s):
            doc.add_heading(s, level=1)
            i += 1
            continue

        # Bullet point
        m = RE_BULLET.match(s)
        if m:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(m.group(1))
            i += 1
            continue

        # Indented lines (4+ spaces) — sub-items / page listings
        leading = len(line) - len(line.lstrip())
        if leading >= 4:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(s)
            set_run_font(run, size=11)
            i += 1
            continue

        # Table block
        if is_table_header(s):
            table_lines = []
            j = i
            while j < n and lines[j].strip():
                table_lines.append(lines[j])
                j += 1
            # Only create table if >= 2 rows and all parseable
            rows_check = [split_cols(l) for l in table_lines if l.strip()]
            if len(rows_check) >= 2 and all(len(r) >= 2 for r in rows_check):
                add_table_block(doc, table_lines)
                i = j
                continue

        # Sub-heading (Heading 4)
        if is_subheading(line, prev_blank, next_blank):
            doc.add_heading(s, level=4)
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        run = p.add_run(s)
        set_run_font(run, size=11)
        i += 1

    doc.save(OUTPUT)
    print(f"OK Guardado: {OUTPUT}")


if __name__ == '__main__':
    main()
